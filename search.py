#!/usr/bin/env python3
"""
Card Cutter - Debate research tool
Searches for evidence using Exa's semantic search with LLM-powered
query generation, filtering, and ranking.
"""

import os
import re
import json
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from dotenv import load_dotenv
from exa_py import Exa
from bs4 import BeautifulSoup
from openai import OpenAI
import trafilatura
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

load_dotenv()


@dataclass
class SearchResult:
    """Normalized search result with cleaned text."""
    title: str
    url: str
    text: str  # Snippet from Exa
    published_date: str = None
    score: float = 0.0  # Ranking score (0-10)
    full_text: str = ""  # Full article text from trafilatura


@dataclass
class Card:
    """A formatted debate card."""
    tag: str  # The argument summary (e.g., "Nuclear war doesn't cause extinction")
    # Citation components
    author_name: str  # "Smith" or "Smith and Jones"
    author_year: str  # "24" (two digit year)
    author_quals: str  # "Professor of Political Science at Harvard"
    article_title: str  # "The End of Trump's Political Capital"
    url: str
    body: str  # The raw excerpt text
    # Markup: list of (start, end, level) tuples where level is 'underline' or 'highlight'
    markup: list = field(default_factory=list)

    @property
    def cite_short(self) -> str:
        """The short cite read aloud: 'Smith 24'"""
        return f"{self.author_name} {self.author_year}"

    @property
    def cite_full(self) -> str:
        """Full formatted cite block."""
        lines = [f"{self.author_name} {self.author_year} – {self.author_quals}"]
        lines.append(f'"{self.article_title}"')
        lines.append(self.url)
        return '\n'.join(lines)

    def render_terminal(self) -> str:
        """Render card for terminal display with ANSI formatting."""
        # Sort markup by start position
        sorted_markup = sorted(self.markup, key=lambda x: x[0])

        # Build the formatted body
        result = []
        pos = 0
        body = self.body

        # Create a character-level format map
        char_format = ['small'] * len(body)  # Default to small
        for start, end, level in sorted_markup:
            for i in range(start, min(end, len(body))):
                if level == 'highlight':
                    char_format[i] = 'highlight'
                elif level == 'underline' and char_format[i] != 'highlight':
                    char_format[i] = 'underline'

        # Render with format changes
        current_format = None
        for i, char in enumerate(body):
            fmt = char_format[i]
            if fmt != current_format:
                # Close previous format
                if current_format == 'highlight':
                    result.append('\033[0m')  # Reset
                elif current_format == 'underline':
                    result.append('\033[0m')
                elif current_format == 'small':
                    result.append('\033[0m')

                # Open new format
                if fmt == 'highlight':
                    result.append('\033[1;43m')  # Bold + yellow background
                elif fmt == 'underline':
                    result.append('\033[4m')  # Underline
                elif fmt == 'small':
                    result.append('\033[2m')  # Dim

                current_format = fmt
            result.append(char)

        # Close final format
        if current_format:
            result.append('\033[0m')

        formatted_body = ''.join(result)

        # Build full card
        output = []
        output.append(f"\033[1;32m{self.tag}\033[0m")  # Green bold tag
        # Cite: bold name+year, then quals, then title, then URL
        output.append(f"\033[1m{self.author_name} {self.author_year}\033[0m – {self.author_quals}")
        output.append(f'"{self.article_title}"')
        output.append(f"\033[2m{self.url}\033[0m")  # Dim URL
        output.append("")
        output.append(formatted_body)

        return '\n'.join(output)

    def render_plain(self) -> str:
        """Render card as plain text with markers."""
        # Sort markup by start position
        sorted_markup = sorted(self.markup, key=lambda x: x[0])

        # Build the formatted body with text markers
        char_format = ['small'] * len(self.body)
        for start, end, level in sorted_markup:
            for i in range(start, min(end, len(self.body))):
                if level == 'highlight':
                    char_format[i] = 'highlight'
                elif level == 'underline' and char_format[i] != 'highlight':
                    char_format[i] = 'underline'

        # Build with markers: CAPS for highlight, regular for underline, (parens) for small
        result = []
        current_format = None
        buffer = []

        for i, char in enumerate(self.body):
            fmt = char_format[i]
            if fmt != current_format:
                # Flush buffer with appropriate formatting
                if buffer:
                    text = ''.join(buffer)
                    if current_format == 'highlight':
                        result.append(text.upper())
                    elif current_format == 'underline':
                        result.append(text)
                    elif current_format == 'small':
                        result.append(f"({text})")
                    buffer = []
                current_format = fmt
            buffer.append(char)

        # Flush remaining buffer
        if buffer:
            text = ''.join(buffer)
            if current_format == 'highlight':
                result.append(text.upper())
            elif current_format == 'underline':
                result.append(text)
            elif current_format == 'small':
                result.append(f"({text})")

        formatted_body = ''.join(result)

        output = []
        output.append(self.tag)
        output.append(f"{self.author_name} {self.author_year} – {self.author_quals}")
        output.append(f'"{self.article_title}"')
        output.append(self.url)
        output.append("")
        output.append(formatted_body)

        return '\n'.join(output)

    def save_to_docx(self, filepath: str) -> str:
        """
        Save card to a Word document with proper debate formatting.

        Formatting:
        - Tag: Bold, 12pt
        - Cite: Bold, 11pt
        - URL: Small, 8pt
        - Body:
          - Highlighted: Bold + Yellow highlight + Underline, 11pt
          - Underlined: Underline, 11pt
          - Small (context): No formatting, 8pt
        """
        doc = Document()

        # Add tag (bold, larger)
        tag_para = doc.add_paragraph()
        tag_run = tag_para.add_run(self.tag)
        tag_run.bold = True
        tag_run.font.size = Pt(12)

        # Add cite line 1: Author Year (bold) – qualifications (not bold)
        cite_para = doc.add_paragraph()
        # Bold part: "Smith 24"
        name_year_run = cite_para.add_run(f"{self.author_name} {self.author_year}")
        name_year_run.bold = True
        name_year_run.font.size = Pt(11)
        # Non-bold part: " – qualifications"
        quals_run = cite_para.add_run(f" – {self.author_quals}")
        quals_run.font.size = Pt(11)

        # Add article title in quotes
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f'"{self.article_title}"')
        title_run.font.size = Pt(10)

        # Add URL (small, gray)
        url_para = doc.add_paragraph()
        url_run = url_para.add_run(self.url)
        url_run.font.size = Pt(8)
        url_run.font.color.rgb = RGBColor(128, 128, 128)

        # Build character-level format map
        char_format = ['small'] * len(self.body)
        sorted_markup = sorted(self.markup, key=lambda x: x[0])

        for start, end, level in sorted_markup:
            for i in range(start, min(end, len(self.body))):
                if level == 'highlight':
                    char_format[i] = 'highlight'
                elif level == 'underline' and char_format[i] != 'highlight':
                    char_format[i] = 'underline'

        # Add body with formatting
        body_para = doc.add_paragraph()

        current_format = None
        buffer = []

        def flush_buffer():
            if not buffer:
                return
            text = ''.join(buffer)
            run = body_para.add_run(text)

            if current_format == 'highlight':
                run.bold = True
                run.underline = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.size = Pt(11)
            elif current_format == 'underline':
                run.underline = True
                run.font.size = Pt(11)
            else:  # small
                run.font.size = Pt(8)

            buffer.clear()

        for i, char in enumerate(self.body):
            fmt = char_format[i]
            if fmt != current_format:
                flush_buffer()
                current_format = fmt
            buffer.append(char)

        # Flush remaining
        flush_buffer()

        # Save document
        doc.save(filepath)
        return filepath

    def add_to_docx(self, doc: Document) -> Document:
        """
        Add this card to an existing Word document.
        Useful for compiling multiple cards into one file.
        """
        # Add a page break if not the first card
        if len(doc.paragraphs) > 0:
            doc.add_paragraph()  # Spacing between cards

        # Add tag (bold, larger)
        tag_para = doc.add_paragraph()
        tag_run = tag_para.add_run(self.tag)
        tag_run.bold = True
        tag_run.font.size = Pt(12)

        # Add cite line 1: Author Year (bold) – qualifications (not bold)
        cite_para = doc.add_paragraph()
        name_year_run = cite_para.add_run(f"{self.author_name} {self.author_year}")
        name_year_run.bold = True
        name_year_run.font.size = Pt(11)
        quals_run = cite_para.add_run(f" – {self.author_quals}")
        quals_run.font.size = Pt(11)

        # Add article title in quotes
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f'"{self.article_title}"')
        title_run.font.size = Pt(10)

        # Add URL (small, gray)
        url_para = doc.add_paragraph()
        url_run = url_para.add_run(self.url)
        url_run.font.size = Pt(8)
        url_run.font.color.rgb = RGBColor(128, 128, 128)

        # Build character-level format map
        char_format = ['small'] * len(self.body)
        sorted_markup = sorted(self.markup, key=lambda x: x[0])

        for start, end, level in sorted_markup:
            for i in range(start, min(end, len(self.body))):
                if level == 'highlight':
                    char_format[i] = 'highlight'
                elif level == 'underline' and char_format[i] != 'highlight':
                    char_format[i] = 'underline'

        # Add body with formatting
        body_para = doc.add_paragraph()

        current_format = None
        buffer = []

        def flush_buffer():
            if not buffer:
                return
            text = ''.join(buffer)
            run = body_para.add_run(text)

            if current_format == 'highlight':
                run.bold = True
                run.underline = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.size = Pt(11)
            elif current_format == 'underline':
                run.underline = True
                run.font.size = Pt(11)
            else:  # small
                run.font.size = Pt(8)

            buffer.clear()

        for i, char in enumerate(self.body):
            fmt = char_format[i]
            if fmt != current_format:
                flush_buffer()
                current_format = fmt
            buffer.append(char)

        flush_buffer()

        return doc


def get_openai_client():
    """Get OpenAI client, raising error if no API key."""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY not set in environment")
    return OpenAI(api_key=api_key)


def get_exa_client():
    """Get Exa client, raising error if no API key."""
    api_key = os.getenv("EXA_API_KEY")
    if not api_key:
        raise ValueError("EXA_API_KEY not set in environment")
    return Exa(api_key=api_key)


def clean_text(text: str) -> str:
    """Clean extracted text by removing HTML artifacts and noise."""
    if not text:
        return ""

    soup = BeautifulSoup(text, "lxml")
    for element in soup(["script", "style", "nav", "header", "footer", "aside"]):
        element.decompose()

    text = soup.get_text(separator=" ")

    noise_patterns = [
        r'\[iframe\][^\]]*',
        r'Advertisement',
        r'Opens in a new window.*?(?=\s{2}|\n|$)',
        r'Close this dialog',
        r'Privacy Policy',
        r'Cookie Policy',
        r'Log\s*in',
        r'Sign\s*up',
        r'Subscribe',
        r'Newsletter',
        r'Share this',
        r'Follow us',
        r'Related Articles?',
        r'Read More',
        r'Click here',
        r'Learn more',
        r'\bMenu\b',
        r'\bSearch\b(?!\w)',
        r'Skip to (?:main )?content',
        r'Accept (?:all )?cookies?',
    ]

    for pattern in noise_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)

    text = re.sub(r'\s+', ' ', text).strip()
    return text


def fetch_full_article(url: str) -> str:
    """
    Fetch and extract full article text from a URL using trafilatura.
    Returns empty string on failure.
    """
    try:
        downloaded = trafilatura.fetch_url(url)
        if downloaded:
            text = trafilatura.extract(
                downloaded,
                include_comments=False,
                include_tables=False,
                no_fallback=False,
            )
            return text or ""
    except Exception as e:
        print(f"    Failed to fetch {url[:40]}...: {e}")
    return ""


def fetch_articles_parallel(results: list[SearchResult], max_workers: int = 4) -> list[SearchResult]:
    """
    Fetch full article text for multiple results in parallel.
    Updates results in place and returns them.
    """
    def fetch_for_result(result: SearchResult) -> tuple[SearchResult, str]:
        text = fetch_full_article(result.url)
        return result, text

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(fetch_for_result, r): r for r in results}

        for future in as_completed(futures):
            try:
                result, full_text = future.result()
                result.full_text = full_text
                if full_text:
                    print(f"    ✓ {result.title[:50]}... ({len(full_text)} chars)")
                else:
                    print(f"    ✗ {result.title[:50]}... (fetch failed)")
            except Exception as e:
                print(f"    Error: {e}")

    return results


def generate_search_queries(user_prompt: str, num_queries: int = 4) -> list[str]:
    """
    Use LLM to generate multiple search queries from user's natural language description.
    """
    client = get_openai_client()

    system_prompt = """You are a search query generator for debate research.
Given a user's description of evidence they're looking for, generate diverse search queries that will find relevant articles.

Generate queries that:
1. Use different phrasings and synonyms
2. Target different aspects of the argument
3. Include some queries with specific terms likely to appear in relevant articles
4. Vary between broad and specific queries

Return a JSON array of query strings, nothing else."""

    user_msg = f"""Generate {num_queries} search queries to find evidence for:

"{user_prompt}"

Return only a JSON array of strings."""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_msg}
        ],
        max_tokens=500,
        temperature=0.7,
    )

    content = response.choices[0].message.content.strip()
    # Handle potential markdown code blocks
    if content.startswith("```"):
        content = re.sub(r'^```\w*\n?', '', content)
        content = re.sub(r'\n?```$', '', content)

    queries = json.loads(content)
    return queries


def run_exa_search(query: str, num_results: int = 10,
                   start_date: str = None, end_date: str = None) -> list[SearchResult]:
    """
    Run a single Exa search and return normalized results.
    """
    exa = get_exa_client()

    kwargs = {
        "type": "neural",
        "num_results": num_results,
        "text": {"max_characters": 3000},
    }

    if start_date:
        kwargs["start_published_date"] = start_date
    if end_date:
        kwargs["end_published_date"] = end_date

    try:
        results = exa.search_and_contents(query, **kwargs)

        normalized = []
        for r in results.results:
            normalized.append(SearchResult(
                title=r.title or "Untitled",
                url=r.url,
                text=clean_text(r.text) if r.text else "",
                published_date=getattr(r, 'published_date', None),
            ))
        return normalized
    except Exception as e:
        print(f"  Error searching '{query[:30]}...': {e}")
        return []


def run_parallel_searches(queries: list[str], results_per_query: int = 10,
                          start_date: str = None, end_date: str = None) -> list[SearchResult]:
    """
    Run multiple Exa searches in parallel and dedupe results by URL.
    """
    all_results = []
    seen_urls = set()

    with ThreadPoolExecutor(max_workers=4) as executor:
        futures = {
            executor.submit(run_exa_search, q, results_per_query, start_date, end_date): q
            for q in queries
        }

        for future in as_completed(futures):
            query = futures[future]
            try:
                results = future.result()
                for r in results:
                    if r.url not in seen_urls and r.text and len(r.text) > 200:
                        seen_urls.add(r.url)
                        all_results.append(r)
            except Exception as e:
                print(f"  Search failed for '{query[:30]}...': {e}")

    return all_results


def batch_filter_results(results: list[SearchResult], user_prompt: str,
                         batch_size: int = 5) -> list[SearchResult]:
    """
    Filter results in batches using LLM, keeping only those supporting the argument.
    """
    if not results:
        return []

    client = get_openai_client()
    filtered = []

    for i in range(0, len(results), batch_size):
        batch = results[i:i + batch_size]

        # Build batch evaluation prompt
        articles = []
        for j, r in enumerate(batch):
            articles.append(f"""ARTICLE {j+1}:
Title: {r.title}
Text: {r.text[:1500]}
---""")

        prompt = f"""You are evaluating whether articles support a specific argument for debate research.

USER'S DESIRED ARGUMENT:
"{user_prompt}"

ARTICLES TO EVALUATE:
{chr(10).join(articles)}

For each article, determine if it provides evidence that SUPPORTS the user's desired argument.
Return a JSON array of objects with "article_num" (1-indexed) and "supports" (true/false).
Only return the JSON array, nothing else.

Example output: [{{"article_num": 1, "supports": true}}, {{"article_num": 2, "supports": false}}]"""

        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=200,
                temperature=0,
            )

            content = response.choices[0].message.content.strip()
            if content.startswith("```"):
                content = re.sub(r'^```\w*\n?', '', content)
                content = re.sub(r'\n?```$', '', content)

            evaluations = json.loads(content)

            for eval_item in evaluations:
                idx = eval_item["article_num"] - 1
                if eval_item["supports"] and 0 <= idx < len(batch):
                    filtered.append(batch[idx])
                    print(f"  ✓ {batch[idx].title[:55]}...")
                elif 0 <= idx < len(batch):
                    print(f"  ✗ {batch[idx].title[:55]}...")

        except Exception as e:
            print(f"  Error in batch filter: {e}")
            # On error, keep all results from this batch
            filtered.extend(batch)

    return filtered


def rank_results(results: list[SearchResult], user_prompt: str) -> list[SearchResult]:
    """
    Rank results by how strongly they support the user's argument (0-10 scale).
    """
    if not results:
        return []

    client = get_openai_client()

    # Build ranking prompt with all results
    articles = []
    for i, r in enumerate(results):
        articles.append(f"""ARTICLE {i+1}:
Title: {r.title}
Text: {r.text[:1000]}
---""")

    prompt = f"""You are ranking debate evidence by how strongly it supports an argument.

USER'S DESIRED ARGUMENT:
"{user_prompt}"

ARTICLES TO RANK:
{chr(10).join(articles)}

Score each article from 0-10 based on how strongly and directly it supports the argument:
- 10: Explicitly and strongly supports the exact argument with clear evidence
- 7-9: Clearly supports the argument with good evidence
- 4-6: Somewhat supports or implies the argument
- 1-3: Weakly or tangentially related
- 0: Does not support

Return a JSON array of objects with "article_num" (1-indexed) and "score" (0-10).
Only return the JSON array."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300,
            temperature=0,
        )

        content = response.choices[0].message.content.strip()
        if content.startswith("```"):
            content = re.sub(r'^```\w*\n?', '', content)
            content = re.sub(r'\n?```$', '', content)

        scores = json.loads(content)

        for score_item in scores:
            idx = score_item["article_num"] - 1
            if 0 <= idx < len(results):
                results[idx].score = score_item["score"]

        # Sort by score descending
        results.sort(key=lambda x: x.score, reverse=True)

    except Exception as e:
        print(f"  Error ranking results: {e}")

    return results


def cut_card(result: SearchResult, user_prompt: str, highlight_level: str = "moderate") -> Card:
    """
    Use LLM to cut a debate card from a search result.

    Args:
        result: SearchResult with full_text
        user_prompt: The user's desired argument
        highlight_level: "conservative", "moderate", or "aggressive"

    Returns:
        A formatted Card object
    """
    client = get_openai_client()

    # Adjust instructions based on highlight level
    highlight_instructions = {
        "conservative": "Highlight 5-7 complete sentences that are the core claims and warrants. Underline an additional 8-12 sentences for context. The highlighted portion when read aloud should take 15-20 seconds.",
        "moderate": "Highlight 7-10 complete sentences that directly support the argument. Underline an additional 5-8 sentences for context. The highlighted portion when read aloud should take 20-30 seconds.",
        "aggressive": "Highlight 10-15 complete sentences - be generous with what gets read. Underline minimal additional context (3-5 sentences). The highlighted portion when read aloud should take 30-45 seconds.",
    }

    prompt = f"""You are an expert policy debate researcher cutting a card from an article.

USER'S DESIRED ARGUMENT:
"{user_prompt}"

ARTICLE TEXT:
\"\"\"
{result.full_text[:8000]}
\"\"\"

ARTICLE METADATA:
- Title: {result.title}
- URL: {result.url}
- Published: {result.published_date or "Unknown"}

Your task:
1. Write a TAG (1 sentence, punchy summary of the argument this evidence proves)
2. Extract CITATION info:
   - author_name: Author's last name (or "Last1 and Last2" for two authors, or publication name if no author)
   - author_year: Two-digit year (e.g., "24" for 2024)
   - author_quals: 1-2 sentences describing why this author is credible (title, affiliation, expertise). Look for bylines, author bios, or "About the author" sections. If unknown, use "Staff writer at [Publication]" or similar.
   - article_title: The title of the article
3. Select the best EXCERPT - THIS IS CRITICAL: Select 3-5 full paragraphs (400-800 words, 15-25 sentences). Debate cards are substantial - they need enough content to make a complete argument with warrants and evidence. Do NOT select just a few sentences.
4. Mark what to UNDERLINE - Most of the excerpt should be underlined (70-90%). Only leave small connecting phrases un-underlined.
5. Mark what to HIGHLIGHT (subset of underlined) - These are complete sentences read aloud at speed.

{highlight_instructions.get(highlight_level, highlight_instructions["moderate"])}

Formatting rules for debate cards:
- Highlighted text = read aloud at speed. Should be COMPLETE SENTENCES, not fragments. A typical card has 5-15 highlighted sentences.
- Underlined text = judge reads but debater doesn't speak (supporting context, secondary warrants)
- Small font (not underlined) = minimal connecting words only, like "He said that" or "According to the report,"

CRITICAL: Debate cards are LONG. A good card has:
- 3-5 paragraphs of text
- 15-25 total sentences in the excerpt
- 5-15 complete sentences highlighted
- Almost everything underlined except brief transitions

Return a JSON object with:
{{
    "tag": "The argument summary",
    "author_name": "Smith",
    "author_year": "24",
    "author_quals": "Professor of Political Science at Harvard, former NSC advisor",
    "article_title": "The End of Trump's Political Capital",
    "excerpt": "The FULL selected passage - 3-5 paragraphs, 400-800 words minimum",
    "underlined": ["sentence 1 to underline", "sentence 2 to underline", ...],
    "highlighted": ["complete sentence 1 to highlight", "complete sentence 2 to highlight", ...]
}}

IMPORTANT:
- The excerpt MUST be 400-800 words (3-5 paragraphs). Short cards are useless in debate.
- The "underlined" array should contain COMPLETE SENTENCES that appear exactly in "excerpt"
- The "highlighted" array should contain 5-15 COMPLETE SENTENCES (not phrases)
- For author_quals, be specific about credentials
- Return only the JSON object, no other text."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4000,  # Increased for longer excerpts
            temperature=0.3,
        )

        content = response.choices[0].message.content.strip()
        if content.startswith("```"):
            content = re.sub(r'^```\w*\n?', '', content)
            content = re.sub(r'\n?```$', '', content)

        data = json.loads(content)

        # Build markup from underlined and highlighted phrases
        markup = []
        excerpt = data["excerpt"]

        # First pass: find all underlined phrases
        for phrase in data.get("underlined", []):
            start = excerpt.find(phrase)
            if start != -1:
                markup.append((start, start + len(phrase), "underline"))

        # Second pass: upgrade highlighted phrases
        for phrase in data.get("highlighted", []):
            start = excerpt.find(phrase)
            if start != -1:
                markup.append((start, start + len(phrase), "highlight"))

        return Card(
            tag=data["tag"],
            author_name=data.get("author_name", "Unknown"),
            author_year=data.get("author_year", "n.d."),
            author_quals=data.get("author_quals", ""),
            article_title=data.get("article_title", result.title),
            url=result.url,
            body=excerpt,
            markup=markup,
        )

    except Exception as e:
        print(f"  Error cutting card: {e}")
        # Extract year from published_date if available
        year = "n.d."
        if result.published_date:
            year = result.published_date[2:4] if len(result.published_date) >= 4 else "n.d."
        # Return a basic card without markup
        return Card(
            tag=f"[Error generating tag: {e}]",
            author_name=result.title.split()[0] if result.title else "Unknown",
            author_year=year,
            author_quals="",
            article_title=result.title or "Unknown",
            url=result.url,
            body=result.full_text[:1000] if result.full_text else result.text[:500],
            markup=[],
        )


def format_result(result: SearchResult, index: int, show_full: bool = False) -> str:
    """Format a single search result for display."""
    output = []
    output.append(f"\n{'='*60}")
    score_str = f" [Score: {result.score}/10]" if result.score > 0 else ""
    output.append(f"[{index}]{score_str} {result.title}")
    output.append(f"URL: {result.url}")
    if result.published_date:
        output.append(f"Published: {result.published_date[:10]}")
    if result.full_text:
        output.append(f"Full article: {len(result.full_text)} chars")

    # Show full text or preview
    text_to_show = result.full_text if (show_full and result.full_text) else result.text
    if text_to_show:
        if show_full:
            output.append(f"\n{text_to_show}")
        else:
            preview = text_to_show[:700].strip()
            if len(text_to_show) > 700:
                preview += "..."
            output.append(f"\n{preview}")
    output.append(f"{'='*60}")
    return "\n".join(output)


def search_pipeline(user_prompt: str, start_date: str = None,
                    num_queries: int = 4, results_per_query: int = 10,
                    fetch_full_text: bool = True) -> list[SearchResult]:
    """
    Full search pipeline:
    1. Generate search queries from user prompt
    2. Run parallel Exa searches
    3. Batch filter by argument relevance
    4. Fetch full article text
    5. Rank by argument strength
    """
    print(f"\n[1/5] Generating search queries...")
    queries = generate_search_queries(user_prompt, num_queries)
    for i, q in enumerate(queries, 1):
        print(f"  {i}. {q}")

    print(f"\n[2/5] Searching Exa ({len(queries)} queries, {results_per_query} results each)...")
    all_results = run_parallel_searches(queries, results_per_query, start_date)
    print(f"  Found {len(all_results)} unique results")

    if not all_results:
        return []

    print(f"\n[3/5] Filtering by argument relevance...")
    filtered = batch_filter_results(all_results, user_prompt)
    print(f"\n  {len(filtered)} results support the argument")

    if not filtered:
        return []

    if fetch_full_text:
        print(f"\n[4/5] Fetching full article text...")
        fetch_articles_parallel(filtered)
        # Filter out results where we couldn't get full text
        filtered = [r for r in filtered if r.full_text]
        print(f"\n  {len(filtered)} articles successfully fetched")

        if not filtered:
            return []

    print(f"\n[5/5] Ranking by argument strength...")
    ranked = rank_results(filtered, user_prompt)

    return ranked


def main():
    print("=" * 60)
    print("CARD CUTTER - Debate Research Tool")
    print("=" * 60)
    print("\nDescribe the card you're looking for in natural language.")
    print("You can specify time constraints like 'from the last year'.")
    print("\nExample: 'Find evidence that Trump has lost political capital")
    print("         and can't get legislation through Congress. Last year.'")
    print("\nEnter your query (or 'quit' to exit):\n")

    while True:
        user_input = input("> ").strip()

        if user_input.lower() in ('quit', 'exit', 'q'):
            print("Goodbye!")
            break

        if not user_input:
            continue

        # Simple date extraction (could be more sophisticated)
        start_date = None
        if "last year" in user_input.lower() or "past year" in user_input.lower():
            start_date = "2024-01-01"
        elif "last month" in user_input.lower():
            start_date = "2024-12-01"
        elif "2024" in user_input:
            start_date = "2024-01-01"
        elif "2025" in user_input:
            start_date = "2025-01-01"

        try:
            results = search_pipeline(user_input, start_date=start_date)

            if not results:
                print("\nNo results found that support your argument.")
                print("Try rephrasing or broadening your search.")
            else:
                print(f"\n{'='*60}")
                print(f"TOP RESULTS ({len(results)} found)")
                print(f"{'='*60}")

                for i, result in enumerate(results[:10], 1):  # Show top 10
                    print(format_result(result, i))

        except Exception as e:
            print(f"Error: {e}")
            import traceback
            traceback.print_exc()

        print("\n" + "-"*60)
        print("Enter another query or 'quit' to exit:\n")


if __name__ == "__main__":
    main()
